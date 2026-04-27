# ─────────────────────────────────────────────
#  pipeline/resume_parser.py
#  Reads PDFs and DOCX, extracts structured features
# ─────────────────────────────────────────────

import os
import re
import sys
import logging
from pathlib import Path

import pdfplumber
from docx import Document

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
from config import (
    SKILL_TAXONOMY, DOMAIN_KEYWORDS,
    EXPERIENCE_PATTERNS, MAX_EXPERIENCE_CAP,
    DOMAIN_PARTIAL_CREDIT
)

logger = logging.getLogger(__name__)

# ── Education section markers ──────────────────
# Year spans near these words are school years, not work years
EDUCATION_MARKERS = [
    "school", "college", "university", "institute", "education",
    "bachelor", "master", "b.tech", "m.tech", "b.e", "m.e",
    "degree", "graduated", "graduation", "cgpa", "gpa", "class of",
    "high school", "secondary", "senior secondary", "12th", "10th",
]

# ── Fresher indicators ─────────────────────────
FRESHER_INDICATORS = [
    "fresher", "fresh graduate", "recent graduate", "final year",
    "final-year", "currently pursuing", "expected graduation",
    "graduating in", "pursuing b", "pursuing m", "currently in",
    "3rd year", "4th year", "second year", "third year", "fourth year",
    "– 2025", "– 2026", "– 2027", "- 2025", "- 2026", "- 2027",
    "2025 – present", "2026 – present", "expected 2025", "expected 2026",
    "batch of 2025", "batch of 2026", "class of 2025", "class of 2026",
]


# ── Text extraction ────────────────────────────

def extract_text_from_pdf(filepath: str) -> str:
    try:
        with pdfplumber.open(filepath) as pdf:
            pages = [page.extract_text() or "" for page in pdf.pages]
        return "\n".join(pages)
    except Exception as e:
        logger.warning(f"PDF parse failed [{filepath}]: {e}")
        return ""


def extract_text_from_docx(filepath: str) -> str:
    try:
        doc = Document(filepath)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n".join(paragraphs)
    except Exception as e:
        logger.warning(f"DOCX parse failed [{filepath}]: {e}")
        return ""


def extract_text(filepath: str) -> str:
    ext = Path(filepath).suffix.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(filepath)
    elif ext in (".docx", ".doc"):
        return extract_text_from_docx(filepath)
    else:
        logger.warning(f"Unsupported file type: {filepath}")
        return ""


# ── Skill extraction ───────────────────────────

def extract_skills(text: str) -> list:
    text_lower = text.lower()
    found = set()
    sorted_terms = sorted(SKILL_TAXONOMY.keys(), key=len, reverse=True)
    for term in sorted_terms:
        pattern = r'\b' + re.escape(term) + r'\b'
        if re.search(pattern, text_lower):
            found.add(SKILL_TAXONOMY[term])
    return sorted(list(found))


# ── Experience extraction ──────────────────────

def _is_near_education(text: str, match_start: int, window: int = 200) -> bool:
    """
    Checks if a year span match is near education-related words.
    Looks at the surrounding text window to decide.
    """
    start = max(0, match_start - window)
    end   = min(len(text), match_start + window)
    surrounding = text[start:end].lower()
    return any(marker in surrounding for marker in EDUCATION_MARKERS)


def _is_fresher(text: str) -> bool:
    """Returns True if the resume shows clear signs of being a student/fresher."""
    text_lower = text.lower()
    return any(indicator in text_lower for indicator in FRESHER_INDICATORS)


def extract_experience(text: str) -> float:
    """
    Extracts years of experience from resume text.
    - First tries explicit patterns like "3 years of experience"
    - Falls back to year spans, but filters out education year spans
    - Detects freshers/students and caps at 1.0 year
    - Returns 0.0 if nothing meaningful found
    """
    text_lower = text.lower()

    # Try explicit experience patterns first
    for pattern in EXPERIENCE_PATTERNS:
        match = re.search(pattern, text_lower)
        if match:
            # Make sure this isn't near education context
            if not _is_near_education(text_lower, match.start()):
                years = float(match.group(1))
                return min(years, MAX_EXPERIENCE_CAP)

    # Fallback: year span detection — but filter education spans
    year_span_pattern = r'\b(20\d{2}|19\d{2})\b.*?\b(20\d{2}|19\d{2})\b'
    matches = list(re.finditer(year_span_pattern, text_lower))

    valid_durations = []
    for m in matches:
        start_yr = int(m.group(1))
        end_yr   = int(m.group(2))
        diff     = end_yr - start_yr

        if diff <= 0 or diff > MAX_EXPERIENCE_CAP:
            continue

        # Skip if this span is near education keywords
        if _is_near_education(text_lower, m.start()):
            continue

        valid_durations.append(diff)

    if valid_durations:
        return min(max(valid_durations), MAX_EXPERIENCE_CAP)

    # If no experience found but clear fresher signals — assign 1.0
    # Freshers have real project experience worth ~1 year equivalent
    if _is_fresher(text_lower):
        return 1.0

    return 0.0


# ── Domain classification ──────────────────────

def classify_domain(text: str, skills: list) -> str:
    """
    Scores each domain based on keyword hits in text + skill overlap.
    Special handling for fullstack profiles — if someone scores high in
    both frontend and backend, classify as fullstack instead of picking one.
    Returns the highest-scoring domain. Falls back to 'general'.
    """
    text_lower = text.lower()
    scores = {domain: 0.0 for domain in DOMAIN_KEYWORDS}

    for domain, keywords in DOMAIN_KEYWORDS.items():
        for kw in keywords:
            if kw in text_lower:
                # Title-level matches (job titles, section headers) get higher weight
                if any(kw in line.lower() for line in text.split('\n')
                       if len(line.strip()) < 60):
                    scores[domain] += 2.0
                else:
                    scores[domain] += 1.0
        # Skill list matches add weight
        for skill in skills:
            if skill in keywords:
                scores[domain] += 1.5

    # ── Fullstack detection ──
    # If someone scores meaningfully in BOTH frontend and backend,
    # they're likely fullstack — don't misclassify as frontend or backend
    frontend_score = scores.get("frontend", 0)
    backend_score  = scores.get("backend", 0)
    fullstack_score = scores.get("fullstack", 0)

    if frontend_score > 2 and backend_score > 2:
        # Both frontend and backend signals present — boost fullstack
        scores["fullstack"] = max(
            fullstack_score,
            (frontend_score + backend_score) * 0.6
        )

    best = max(scores, key=scores.get)
    return best if scores[best] > 0 else "general"


# ── Main parser ────────────────────────────────

def parse_resume(filepath: str) -> dict:
    text = extract_text(filepath)
    if not text or len(text.strip()) < 50:
        logger.warning(f"Insufficient text extracted: {filepath}")
        return None

    skills     = extract_skills(text) or []
    experience = extract_experience(text)
    domain     = classify_domain(text, skills)
    resume_id  = Path(filepath).stem

    # Extract candidate name — first short non-empty line heuristic
    name = "Unknown Candidate"
    for line in text.split("\n"):
        if line.strip() and len(line.strip()) < 60:
            name = line.strip()
            break

    return {
        "name":               name,
        "skills":             skills,
        "experience":         experience,
        "domain":             domain,
        "skill_count":        len(skills),
        "raw_text":           text[:3000],
    }